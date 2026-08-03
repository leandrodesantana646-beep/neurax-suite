import streamlit as st
from groq import Groq
import hashlib
from datetime import datetime
import pandas as pd
from supabase import create_client, Client
import pypdf
import io
import random
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from email.mime.application import MIMEApplication
from docx import Document
from fpdf import FPDF

# Configuração inicial da página
st.set_page_config(
    page_title="NeuraX Suite Pro - Life OS",
    page_icon="🚀",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Inicialização do Cliente Supabase via Secrets
try:
    SUPABASE_URL = st.secrets["supabase"]["url"]
    SUPABASE_KEY = st.secrets["supabase"]["key"]
    supabase: Client = create_client(SUPABASE_URL, SUPABASE_KEY)
except Exception as e:
    supabase = None

# Gerenciamento de Sessão Inicial
for key in ["logged_in", "username", "generation_count", "chat_messages", "theme"]:
    if key not in st.session_state:
        if key == "logged_in":
            st.session_state[key] = False
        elif key == "chat_messages":
            st.session_state[key] = []
        elif key == "generation_count":
            st.session_state[key] = 0
        elif key == "theme":
            st.session_state[key] = "Escuro (Dark)"
        else:
            st.session_state[key] = ""

# Seletor de Tema Dinâmico (Claro / Escuro)
with st.sidebar:
    st.title("⚡ NeuraX")
    st.session_state["theme"] = st.selectbox("🎨 Tema da Interface", ["Escuro (Dark)", "Claro (Light)"])

is_dark = "Escuro" in st.session_state["theme"]

# Paleta de Cores baseada no Tema
bg_app = "#0b0f19" if is_dark else "#f8fafc"
text_app = "#f3f4f6" if is_dark else "#1e293b"
sidebar_bg = "#111827" if is_dark else "#ffffff"
card_bg = "#111827" if is_dark else "#ffffff"
border_color = "#1f2937" if is_dark else "#e2e8f0"
chat_input_bottom = "95px"

# Estilização visual customizada com suporte a temas
st.markdown(f"""
    <style>
    @import url('https://fonts.googleapis.com/css2?family=Plus+Jakarta+Sans:wght@400;500;600;700&display=swap');
    
    html, body, [class*="css"] {{
        font-family: 'Plus Jakarta Sans', sans-serif;
    }}
    
    .stApp {{ background-color: {bg_app}; color: {text_app}; }}
    
    [data-testid="stChatInput"] {{
        bottom: {chat_input_bottom} !important;
    }}
    
    footer {{visibility: hidden;}}
    h1, h2, h3 {{ color: #38bdf8 !important; font-weight: 700; }}
    
    [data-testid="stSidebar"] {{
        background-color: {sidebar_bg};
        border-right: 1px solid {border_color};
        padding-top: 1rem;
    }}
    
    [data-testid="stSidebar"] .stRadio > label {{
        font-size: 1.1rem; font-weight: 700; color: #38bdf8 !important; margin-bottom: 12px; letter-spacing: 0.5px;
    }}
    [data-testid="stSidebar"] .stRadio div[role="radiogroup"] {{ gap: 8px; }}
    [data-testid="stSidebar"] .stRadio div[role="radiogroup"] label {{
        background-color: {"#1a2234" if is_dark else "#f1f5f9"}; border: 1px solid {border_color}; padding: 10px 14px; border-radius: 10px;
        transition: all 0.3s ease; cursor: pointer; width: 100%;
    }}
    [data-testid="stSidebar"] .stRadio div[role="radiogroup"] label p,
    [data-testid="stSidebar"] .stRadio div[role="radiogroup"] label span {{
        color: {"#38bdf8" if is_dark else "#0284c7"} !important;
        font-weight: 600;
    }}
    [data-testid="stSidebar"] .stRadio div[role="radiogroup"] label:hover {{
        background-color: #0284c7; border-color: #38bdf8; transform: translateX(4px); box-shadow: 0 4px 12px rgba(56, 189, 248, 0.2);
    }}
    [data-testid="stSidebar"] .stRadio div[role="radiogroup"] label:hover * {{
        color: #ffffff !important;
    }}
    
    .stButton>button {{
        background: linear-gradient(135deg, #0284c7 0%, #38bdf8 100%); color: white; border-radius: 10px;
        font-weight: 600; padding: 0.6rem 1.2rem; border: none; box-shadow: 0 4px 15px rgba(56, 189, 248, 0.3); transition: all 0.3s ease;
    }}
    .stButton>button:hover {{
        background: linear-gradient(135deg, #0369a1 0%, #0284c7 100%); box-shadow: 0 6px 20px rgba(56, 189, 248, 0.5); transform: translateY(-1px);
    }}
    
    [data-testid="stMetric"] {{ background: {card_bg}; border: 1px solid {border_color}; padding: 15px; border-radius: 12px; }}
    .streamlit-expanderHeader {{ background-color: {card_bg}; border: 1px solid {border_color}; border-radius: 8px; }}
    </style>
""", unsafe_allow_html=True)

# Funções de Criptografia, Banco de Dados e Exportação
def make_hash(password):
    return hashlib.sha256(str.encode(password)).hexdigest()

def check_hash(password, hashed_text):
    return make_hash(password) == hashed_text

def add_user(username, password):
    if not supabase:
        st.error("Erro: Supabase não inicializado.")
        return
    try:
        if supabase.table("users").select("username").eq("username", username).execute().data:
            st.warning("Este usuário já existe.")
        else:
            supabase.table("users").insert({"username": username, "password": make_hash(password)}).execute()
            st.success("Cadastro realizado com sucesso! Faça seu login.")
    except Exception as e:
        st.error(f"Erro ao cadastrar: {e}")

def login_user(username, password):
    if username.strip().lower() == "admin" and password.strip().lower() == "admin":
        return True
    if not supabase:
        return False
    try:
        data = supabase.table("users").select("password").eq("username", username).execute().data
        if data and check_hash(password, data[0]["password"]):
            return True
    except Exception:
        pass
    return False

def update_password(username, new_password):
    if not supabase:
        return False
    try:
        hashed = make_hash(new_password)
        supabase.table("users").update({"password": hashed}).eq("username", username).execute()
        return True
    except Exception:
        return False

def send_email_smtp(to_email, subject, body, attachment_bytes=None, attachment_name="relatorio.pdf"):
    try:
        smtp_server = st.secrets.get("smtp", {}).get("server", "smtp.gmail.com")
        smtp_port = int(st.secrets.get("smtp", {}).get("port", 587))
        sender_email = st.secrets.get("smtp", {}).get("email", "seu-email@gmail.com")
        sender_password = st.secrets.get("smtp", {}).get("password", "sua-senha-de-app")
        
        msg = MIMEMultipart()
        msg["From"] = sender_email
        msg["To"] = to_email
        msg["Subject"] = subject
        
        msg.attach(MIMEText(body, "plain"))
        
        if attachment_bytes:
            part = MIMEApplication(attachment_bytes, Name=attachment_name)
            part['Content-Disposition'] = f'attachment; filename="{attachment_name}"'
            msg.attach(part)
            
        server = smtplib.SMTP(smtp_server, smtp_port)
        server.starttls()
        server.login(sender_email, sender_password)
        server.sendmail(sender_email, to_email, msg.as_string())
        server.quit()
        return True
    except Exception as e:
        return False

def save_history(username, tool_name, content, tokens=0):
    if not supabase:
        return
    try:
        timestamp = datetime.now().strftime("%d/%m/%Y %H:%M")
        supabase.table("history").insert({
            "username": username,
            "tool_name": tool_name,
            "content": content,
            "timestamp": timestamp,
            "tokens": tokens
        }).execute()
    except Exception:
        pass

def get_history(username):
    if not supabase:
        return []
    try:
        data = supabase.table("history").select("id, tool_name, content, timestamp, tokens").eq("username", username).order("id", desc=True).execute().data
        return [(r["id"], r["tool_name"], r["content"], r["timestamp"], r.get("tokens", 0)) for r in data]
    except Exception:
        return []

def delete_history_item(item_id):
    if not supabase:
        return
    try:
        supabase.table("history").delete().eq("id", item_id).execute()
    except Exception:
        pass

def get_all_users():
    if not supabase:
        return []
    try:
        data = supabase.table("users").select("username").execute().data
        return [r["username"] for r in data]
    except Exception:
        return []

def get_all_history_admin():
    if not supabase:
        return []
    try:
        data = supabase.table("history").select("username, tool_name, timestamp, tokens").order("id", desc=True).execute().data
        return [(r["username"], r["tool_name"], r["timestamp"], r.get("tokens", 0)) for r in data]
    except Exception:
        return []

def get_user_profile(username):
    if not supabase:
        return {}
    try:
        data = supabase.table("profiles").select("*").eq("username", username).execute().data
        if data:
            return data[0]
    except Exception:
        pass
    return {"business_name": "", "niche": "", "budget": "", "goals": ""}

def save_user_profile(username, business_name, niche, budget, goals):
    if not supabase:
        return
    try:
        existing = supabase.table("profiles").select("username").eq("username", username).execute().data
        if existing:
            supabase.table("profiles").update({
                "business_name": business_name,
                "niche": niche,
                "budget": budget,
                "goals": goals
            }).eq("username", username).execute()
        else:
            supabase.table("profiles").insert({
                "username": username,
                "business_name": business_name,
                "niche": niche,
                "budget": budget,
                "goals": goals
            }).execute()
        st.success("Perfil salvo com sucesso na Memória de Longo Prazo!")
    except Exception as e:
        st.error(f"Erro ao salvar perfil: {e}")

# Funções de Exportação Avançada (Word e PDF)
def export_to_docx(content):
    doc = Document()
    doc.add_heading("NeuraX Suite Pro - Relatório", 0)
    for paragraph in content.split("\n"):
        if paragraph.strip():
            doc.add_paragraph(paragraph)
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

def export_to_pdf(content):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=11)
    clean_content = content.encode('latin-1', 'ignore').decode('latin-1')
    for line in clean_content.split('\n'):
        pdf.multi_cell(0, 8, txt=line)
    return pdf.output(dest='S').encode('latin-1')

# Tela de Autenticação
if not st.session_state["logged_in"]:
    st.title("🚀 NeuraX Suite Pro")
    st.markdown("### O Único Ecossistema Inteligente que Você Precisa")
    auth_mode = st.selectbox("Escolha a opção", ["Login", "Cadastrar", "Esqueci minha senha"])
    
    if auth_mode == "Login":
        user = st.text_input("Usuário ou E-mail")
        pwd = st.text_input("Senha", type="password")
        if st.button("Entrar no Sistema"):
            if login_user(user, pwd):
                st.session_state["logged_in"] = True
                st.session_state["username"] = user
                st.rerun()
            else:
                st.error("Dados incorretos.")
                
    elif auth_mode == "Cadastrar":
        user = st.text_input("E-mail / Usuário")
        pwd = st.text_input("Senha", type="password")
        if st.button("Criar Conta"):
            if user and pwd:
                add_user(user, pwd)
            else:
                st.warning("Preencha todos os campos.")
                
    else:
        st.markdown("### 🔑 Recuperação de Senha")
        if "reset_stage" not in st.session_state:
            st.session_state["reset_stage"] = 1
        
        if st.session_state["reset_stage"] == 1:
            recovery_user = st.text_input("Digite seu E-mail / Usuário cadastrado")
            if st.button("Enviar Código de Verificação"):
                if recovery_user:
                    code = str(random.randint(100000, 999999))
                    st.session_state["reset_code"] = code
                    st.session_state["reset_user"] = recovery_user
                    
                    email_sent = send_email_smtp(recovery_user, "Código de Recuperação - NeuraX", f"Seu código de recuperação é: {code}")
                    st.session_state["email_status"] = "sent" if email_sent else "failed"
                    st.session_state["reset_stage"] = 2
                    st.rerun()
                else:
                    st.warning("Insira seu usuário/e-mail.")
        
        elif st.session_state["reset_stage"] == 2:
            st.info(f"Usuário selecionado: **{st.session_state.get('reset_user')}**")
            if st.session_state.get("email_status") == "sent":
                st.success("📨 Código enviado para o seu e-mail!")
            else:
                st.warning(f"⚠️ **Modo de Teste:** O e-mail não pôde ser enviado por SMTP. Código de verificação: **{st.session_state.get('reset_code')}**")
            
            entered_code = st.text_input("Digite o Código de 6 Dígitos")
            new_pwd = st.text_input("Nova Senha", type="password")
            confirm_pwd = st.text_input("Confirme a Nova Senha", type="password")
            
            col_r1, col_r2 = st.columns(2)
            with col_r1:
                if st.button("Redefinir Senha"):
                    if entered_code == st.session_state.get("reset_code"):
                        if new_pwd and new_pwd == confirm_pwd:
                            target_user = st.session_state.get("reset_user")
                            if update_password(target_user, new_pwd):
                                st.success("Senha redefinida com sucesso! Volte ao Login.")
                                st.session_state["reset_stage"] = 1
                                st.rerun()
                            else:
                                st.error("Erro ao atualizar a senha no banco de dados.")
                        else:
                            st.warning("As senhas não coincidem ou estão vazias.")
                    else:
                        st.error("Código incorreto.")
            with col_r2:
                if st.button("Cancelar"):
                    st.session_state["reset_stage"] = 1
                    st.rerun()

else:
    # Sidebar Principal
    st.sidebar.write(f"Usuário: **{st.session_state['username']}**")
    
    st.sidebar.subheader("⚙️ Preferências de IA")
    user_tone = st.sidebar.selectbox(
        "Tom de Voz", 
        ["Persuasivo & Direto", "Técnico & Profissional", "Divertido & Descontraído", "Empático & Acolhedor"]
    )
    
    model_choice = st.sidebar.selectbox("Modelo", ["Llama-3.3-70b-versatile", "Llama-3.1-8b-instant"])
    model_name = "llama-3.3-70b-versatile" if "70b" in model_choice else "llama-3.1-8b-instant"
    
    groq_api_key = st.secrets.get("GROQ_API_KEY", "")
    if not groq_api_key or not groq_api_key.startswith("gsk_"):
        groq_api_key = st.sidebar.text_input("Groq API Key", type="password")

    client = Groq(api_key=groq_api_key) if groq_api_key else None

    current_profile = get_user_profile(st.session_state["username"])
    profile_context = (
        f"\n[DADOS DE CONTEXTO DO USUÁRIO]:\n"
        f"- Nome do Negócio/Projeto: {current_profile.get('business_name', 'Não informado')}\n"
        f"- Nicho de Atuação: {current_profile.get('niche', 'Não informado')}\n"
        f"- Orçamento Disponível: {current_profile.get('budget', 'Não informado')}\n"
        f"- Objetivos Principais: {current_profile.get('goals', 'Não informado')}\n"
    )

    menu_options = [
        "📊 Meu Painel de Produtividade",
        "💬 Chat Geral com o NeuraX",
        "👤 Meu Perfil & Contexto (Memória)",
        "📂 Analista de Arquivos & PDFs",
        "🏛️ Conselho de Gigantes (Board de IAs)",
        "⚡ Gestor de Tarefas Inteligente",
        "🧠 Mentor de Saúde Mental",
        "📚 Tutor Universal & Estudos",
        "🗺️ Arquiteto de Funis de Vendas",
        "💰 Precificação Inteligente",
        "🎯 Gerador de Anúncios (Meta/Google)",
        "🚀 NeuraX Growth Engine",
        "💸 Gerador de Renda Extra & Negócios",
        "💬 Gerador de Copy WhatsApp",
        "📸 Planejador Instagram",
        "✉️ Gerador de E-mail Comercial",
        "🎬 Gerador de Roteiro para Vídeos",
        "⚖️ Assistente de Burocracias",
        "💸 Consultor de Finanças Pessoais",
        "🍳 Assistente de Despensa & Rotina",
        "🎓 Simulador de Entrevistas",
        "📂 Meu Histórico"
    ]
    
    if st.session_state["username"].strip().lower() == "admin":
        menu_options.insert(0, "🛠️ Painel Administrativo")

    escolha = st.sidebar.radio("⚡ Menu de Ferramentas", menu_options)
    
    if st.sidebar.button("Sair da Conta"):
        st.session_state["logged_in"] = False
        st.session_state["username"] = ""
        st.rerun()

    if escolha == "🛠️ Painel Administrativo":
        st.header("🛠️ Painel Administrativo - NeuraX Suite")
        st.success("Acesso Master Confirmado!")
        all_users = get_all_users()
        all_history = get_all_history_admin()
        
        total_tokens = sum([item[3] for item in all_history])
        estimated_cost = (total_tokens / 1000000) * 0.70
        
        col1, col2, col3 = st.columns(3)
        with col1:
            st.metric("Total de Usuários", len(all_users))
        with col2:
            st.metric("Total de Gerações", len(all_history))
        with col3:
            st.metric("Tokens / Custo Est.", f"{total_tokens:,} (~${estimated_cost:.4f})")
            
        st.markdown("---")
        st.markdown("### 👥 Usuários Registrados")
        st.write(all_users)
        
        st.markdown("### 📊 Histórico Geral de Atividades & Consumo")
        if all_history:
            df_history = pd.DataFrame(all_history, columns=["Usuário", "Ferramenta", "Data/Hora", "Tokens"])
            st.dataframe(df_history, use_container_width=True)
        else:
            st.info("Nenhuma atividade registrada na plataforma ainda.")

    elif escolha == "📊 Meu Painel de Produtividade":
        st.header(f"📊 Painel de Produtividade de {st.session_state['username']}")
        user_history = get_history(st.session_state["username"])
        total_tokens_user = sum([item[4] for item in user_history])
        
        col_m1, col_m2, col_m3 = st.columns(3)
        with col_m1:
            st.metric("Conteúdos Gerados", len(user_history))
        with col_m2:
            st.metric("Ações na Sessão", st.session_state["generation_count"])
        with col_m3:
            st.metric("Tokens Consumidos", f"{total_tokens_user:,}")
        
        if user_history:
            tools_list = [item[1] for item in user_history]
            df_tools = pd.DataFrame(tools_list, columns=["Ferramenta"]).value_counts().reset_index()
            df_tools.columns = ["Ferramenta", "Quantidade"]
            st.markdown("### 📈 Utilização por Ferramenta")
            st.bar_chart(df_tools.set_index("Ferramenta"))
        else:
            st.info("Você ainda não utilizou nenhuma ferramenta.")

    elif escolha == "💬 Chat Geral com o NeuraX":
        st.header("💬 Chat Geral com o NeuraX (Com Entrada por Voz)")
        st.write("Converse digitando ou grave sua voz usando o microfone abaixo.")
        
        for message in st.session_state["chat_messages"]:
            with st.chat_message(message["role"]):
                st.markdown(message["content"])
                
        audio_file = st.audio_input("🎙️ Falar com o NeuraX")
        voice_query = None
        if audio_file and client:
            with st.spinner("Transcrevendo seu áudio com Whisper..."):
                try:
                    transcript_res = client.audio.transcriptions.create(
                        file=("audio.wav", audio_file.getvalue()),
                        model="whisper-large-v3-turbo",
                        response_format="text"
                    )
                    voice_query = transcript_res
                    st.info(f"Transcrição: {voice_query}")
                except Exception as e:
                    st.error(f"Erro ao transcrever áudio: {e}")

        user_query = st.chat_input("Digite sua mensagem para o NeuraX...")
        final_query = voice_query if voice_query else user_query
        
        if final_query:
            if not client:
                st.error("Configure sua chave da Groq no menu lateral.")
            else:
                st.session_state["chat_messages"].append({"role": "user", "content": final_query})
                with st.chat_message("user"):
                    st.markdown(final_query)
                    
                with st.chat_message("assistant"):
                    with st.spinner("Pensando..."):
                        system_prompt = (
                            f"Atue como um Assistente Virtual Inteligente aplicando o tom de voz: '{user_tone}'.\n"
                            f"{profile_context}"
                        )
                        messages_payload = [{"role": "system", "content": system_prompt}] + st.session_state["chat_messages"]
                        completion = client.chat.completions.create(model=model_name, messages=messages_payload)
                        response_text = completion.choices[0].message.content
                        
                        approx_tokens = len(final_query.split()) + len(response_text.split())
                        
                        st.markdown(response_text)
                        st.session_state["chat_messages"].append({"role": "assistant", "content": response_text})
                        st.session_state["generation_count"] += 1
                        save_history(st.session_state["username"], "Chat Geral", response_text, tokens=approx_tokens)

    elif escolha == "👤 Meu Perfil & Contexto (Memória)":
        st.header("👤 Memória de Longo Prazo - Perfil do Usuário")
        with st.form("form_profile"):
            b_name = st.text_input("Nome do Negócio ou Projeto", value=current_profile.get("business_name", ""))
            b_niche = st.text_input("Nicho de Atuação", value=current_profile.get("niche", ""))
            b_budget = st.text_input("Orçamento Disponível / Faturamento", value=current_profile.get("budget", ""))
            b_goals = st.text_area("Objetivos Principais e Metas", value=current_profile.get("goals", ""))
            
            submit_profile = st.form_submit_button("Salvar Perfil Definitivo")
            if submit_profile:
                save_user_profile(st.session_state["username"], b_name, b_niche, b_budget, b_goals)

    elif escolha == "📂 Analista de Arquivos & PDFs":
        st.header("📂 Analista de Arquivos e PDFs")
        arquivo_pdf = st.file_uploader("Envie seu arquivo PDF", type=["pdf"])
        pergunta = st.text_input("O que você deseja saber ou resumir sobre este documento?")
        
        if st.button("Analisar PDF"):
            if not arquivo_pdf or not pergunta or not client:
                st.warning("⚠️ Envie o PDF, digite a pergunta e configure sua chave da Groq.")
            else:
                with st.spinner("Lendo documento..."):
                    try:
                        leitor = pypdf.PdfReader(io.BytesIO(arquivo_pdf.getvalue()))
                        texto_extraido = "".join([p.extract_text() for p in leitor.pages if p.extract_text()])
                        texto_limitado = texto_extraido[:30000]
                        prompt = f"Atue como Especialista (Tom: '{user_tone}').\n{profile_context}\nTexto do PDF:\n{texto_limitado}\nPergunta: {pergunta}"
                        
                        completion = client.chat.completions.create(model=model_name, messages=[{"role": "user", "content": prompt}])
                        resultado = completion.choices[0].message.content
                        approx_tokens = len(prompt.split()) + len(resultado.split())
                        
                        st.session_state["generation_count"] += 1
                        save_history(st.session_state["username"], "Analista de PDF", resultado, tokens=approx_tokens)
                        
                        st.success("Análise Concluída!")
                        st.markdown(resultado)
                        
                        col_ex1, col_ex2 = st.columns(2)
                        with col_ex1:
                            st.download_button("📥 Baixar Word (.docx)", data=export_to_docx(resultado), file_name="analise.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                        with col_ex2:
                            st.download_button("📥 Baixar PDF (.pdf)", data=export_to_pdf(resultado), file_name="analise.pdf", mime="application/pdf")
                    except Exception as e:
                        st.error(f"Erro: {e}")

    elif escolha == "🏛️ Conselho de Gigantes (Board de IAs)":
        st.header("🏛️ Conselho de Gigantes - Board de IAs de Elite")
        st.write("Submeta um desafio crítico e veja 4 lendas de negócios (O Visionário, O Estrategista Financeiro, O Mestre de Growth e o Crítico Implacável) debatendo o seu problema com base no seu perfil cadastrado.")
        
        desafio_board = st.text_area("Qual é o seu dilema ou projeto atual para ser julgado pelo Conselho?")
        
        if st.button("Convocar o Conselho de Gigantes"):
            if not desafio_board or not client:
                st.warning("Insira o seu desafio e verifique sua chave da Groq.")
            else:
                with st.spinner("Convocando os conselheiros e cruzando dados do seu perfil..."):
                    prompt = (
                        f"Atue simultaneamente como um Conselho de Administração de Elite composto por 4 personas brilhantes:\n"
                        f"1. O Visionario (focado em disrupção, inovação e escala massiva)\n"
                        f"2. O Estrategista Financeiro (focado em fluxo de caixa, risco zero e margem de lucro)\n"
                        f"3. O Mestre de Growth (focado em aquisição agressiva, tráfego e conversão)\n"
                        f"4. O Crítico Implacável (Advogado do Diabo, focado em achar falhas e pontos cegos)\n\n"
                        f"{profile_context}\n"
                        f"Desafio submetido pelo usuário: {desafio_board}\n\n"
                        f"Estruture a resposta exatamente assim:\n"
                        f"- **Parecer do Visionário**\n"
                        f"- **Parecer do Estrategista Financeiro**\n"
                        f"- **Parecer do Mestre de Growth**\n"
                        f"- **Alerta do Crítico Implacável**\n"
                        f"- **Consenso do Conselho & Plano de Ação Imediato**"
                    )
                    completion = client.chat.completions.create(model=model_name, messages=[{"role": "user", "content": prompt}])
                    resultado = completion.choices[0].message.content
                    approx_tokens = len(prompt.split()) + len(resultado.split())
                    
                    st.session_state["generation_count"] += 1
                    save_history(st.session_state["username"], "Conselho de Gigantes", resultado, tokens=approx_tokens)
                    st.markdown(resultado)
                    
                    col_ex1, col_ex2 = st.columns(2)
                    with col_ex1:
                        st.download_button("📥 Baixar Word (.docx)", data=export_to_docx(resultado), file_name="conselho_gigantes.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                    with col_ex2:
                        st.download_button("📥 Baixar PDF (.pdf)", data=export_to_pdf(resultado), file_name="conselho_gigantes.pdf", mime="application/pdf")
                        
                    st.markdown("---")
                    st.markdown("### 📧 Enviar Relatório por E-mail")
                    destinatario_email = st.text_input("E-mail de destino", value=st.session_state["username"] if "@" in st.session_state["username"] else "")
                    if st.button("Enviar PDF por E-mail agora"):
                        if destinatario_email:
                            pdf_bytes = export_to_pdf(resultado)
                            enviado = send_email_smtp(destinatario_email, "NeuraX Suite - Conselho de Gigantes", "Olá,\n\nSegue em anexo a ata da reunião do Conselho de Gigantes gerada pelo NeuraX Suite Pro.", pdf_bytes, "conselho_gigantes.pdf")
                            if enviado:
                                st.success("E-mail enviado com sucesso!")
                            else:
                                st.error("Erro ao enviar e-mail. Verifique as credenciais SMTP no st.secrets.")
                        else:
                            st.warning("Insira um endereço de e-mail válido.")

    elif escolha == "💸 Gerador de Renda Extra & Negócios":
        st.header("💸 Gerador de Renda Extra & Negócios Digitais")
        st.write("Descubra caminhos práticos para monetizar suas habilidades e gerar novas fontes de receita.")
        nicho_foco = st.text_input("Qual é a sua principal habilidade ou área de interesse?", value=current_profile.get("niche", ""))
        capital_inicial = st.selectbox("Quanto você tem disponível para investir?", ["Zero (Começar do zero)", "Baixo (Até R$ 500)", "Médio (R$ 500 a R$ 2.000)", "Alto (Acima de R$ 2.000)"])
        
        if st.button("Gerar Plano de Renda Extra"):
            if not client:
                st.warning("Configure sua chave da Groq no menu lateral.")
            else:
                with st.spinner("Mapeando oportunidades de faturamento..."):
                    prompt = (
                        f"Atue como um Especialista em Monetização e Negócios Digitais (Tom: '{user_tone}').\n"
                        f"{profile_context}\n"
                        f"O usuário quer gerar renda extra na área de '{nicho_foco}' com capital inicial '{capital_inicial}'. "
                        f"Crie um plano prático com 3 modelos de negócios viáveis (ex: prestação de serviços, infoprodutos, afiliados), "
                        f"passo a passo para os primeiros 30 dias e sugestões de captação de clientes."
                    )
                    completion = client.chat.completions.create(model=model_name, messages=[{"role": "user", "content": prompt}])
                    resultado = completion.choices[0].message.content
                    approx_tokens = len(prompt.split()) + len(resultado.split())
                    
                    st.session_state["generation_count"] += 1
                    save_history(st.session_state["username"], "Gerador de Renda Extra", resultado, tokens=approx_tokens)
                    st.markdown(resultado)
                    
                    col_ex1, col_ex2 = st.columns(2)
                    with col_ex1:
                        st.download_button("📥 Baixar Word (.docx)", data=export_to_docx(resultado), file_name="renda_extra.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                    with col_ex2:
                        st.download_button("📥 Baixar PDF (.pdf)", data=export_to_pdf(resultado), file_name="renda_extra.pdf", mime="application/pdf")

    elif escolha in [
        "⚡ Gestor de Tarefas Inteligente", "🧠 Mentor de Saúde Mental", "📚 Tutor Universal & Estudos",
        "🗺️ Arquiteto de Funis de Vendas", "💰 Precificação Inteligente", "🎯 Gerador de Anúncios (Meta/Google)",
        "🚀 NeuraX Growth Engine", "💬 Gerador de Copy WhatsApp", "📸 Planejador Instagram", 
        "✉️ Gerador de E-mail Comercial", "🎬 Gerador de Roteiro para Vídeos", "⚖️ Assistente de Burocracias", 
        "💸 Consultor de Finanças Pessoais", "🍳 Assistente de Despensa & Rotina", "🎓 Simulador de Entrevistas"
    ]:
        st.header(escolha)
        detalhe = st.text_area("Descreva os detalhes da sua demanda:")
        if st.button("Gerar com IA"):
            if detalhe and client:
                with st.spinner("Processando..."):
                    prompt = f"Atue como um Especialista (Tom: '{user_tone}').\n{profile_context}\nDemanda: {detalhe}"
                    completion = client.chat.completions.create(model=model_name, messages=[{"role": "user", "content": prompt}])
                    resultado = completion.choices[0].message.content
                    approx_tokens = len(prompt.split()) + len(resultado.split())
                    
                    st.session_state["generation_count"] += 1
                    save_history(st.session_state["username"], escolha, resultado, tokens=approx_tokens)
                    st.markdown(resultado)
                    
                    col_ex1, col_ex2 = st.columns(2)
                    with col_ex1:
                        st.download_button("📥 Baixar Word (.docx)", data=export_to_docx(resultado), file_name="relatorio.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                    with col_ex2:
                        st.download_button("📥 Baixar PDF (.pdf)", data=export_to_pdf(resultado), file_name="relatorio.pdf", mime="application/pdf")
                        
                    st.markdown("---")
                    st.markdown("### 📧 Enviar Relatório por E-mail")
                    destinatario_email = st.text_input("E-mail de destino", value=st.session_state["username"] if "@" in st.session_state["username"] else "")
                    if st.button("Enviar PDF por E-mail agora"):
                        if destinatario_email:
                            pdf_bytes = export_to_pdf(resultado)
                            enviado = send_email_smtp(destinatario_email, f"NeuraX Suite - {escolha}", "Olá,\n\nSegue em anexo o relatório gerado pelo NeuraX Suite Pro.", pdf_bytes, "relatorio.pdf")
                            if enviado:
                                st.success("E-mail enviado com sucesso!")
                            else:
                                st.error("Erro ao enviar e-mail. Verifique as credenciais SMTP no st.secrets.")
                        else:
                            st.warning("Insira um endereço de e-mail válido.")

    elif escolha == "📂 Meu Histórico":
        st.header("📂 Histórico de Gerações & Gerenciamento")
        user_history = get_history(st.session_state["username"])
        
        if not user_history:
            st.info("Nenhum histórico encontrado.")
        else:
            ferramentas_disponiveis = ["Todas"] + list(set([item[1] for item in user_history]))
            filtro_ferramenta = st.selectbox("Filtrar por Ferramenta", ferramentas_disponiveis)
            
            filtered_history = user_history if filtro_ferramenta == "Todas" else [item for item in user_history if item[1] == filtro_ferramenta]
            
            for item_id, tool, content, timestamp, tokens in filtered_history:
                with st.expander(f"🛠️ [{tool}] - {timestamp} (Tokens: {tokens})"):
                    st.markdown(content)
                    
                    col_h1, col_h2, col_h3 = st.columns(3)
                    with col_h1:
                        st.download_button("📥 Baixar (.txt)", data=content, file_name=f"hist_{item_id}.txt", mime="text/plain", key=f"txt_{item_id}")
                    with col_h2:
                        st.download_button("📥 Baixar Word", data=export_to_docx(content), file_name=f"hist_{item_id}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", key=f"docx_{item_id}")
                    with col_h3:
                        if st.button("🗑️ Excluir", key=f"del_{item_id}"):
                            delete_history_item(item_id)
                            st.success("Excluído!")
                            st.rerun()
